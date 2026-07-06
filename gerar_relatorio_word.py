# -*- coding: utf-8 -*-
"""
Módulo para geração do Relatório Executivo em formato Word (.docx).
Lê a planilha Geral.xlsx e monta capa, visão geral de estado (com tabelas e gráficos)
e o detalhamento por unidade e subunidade, conforme regras de negócio.
"""
import io
import os
import re
import tempfile
from datetime import datetime
import pandas as pd

import matplotlib
matplotlib.use("Agg")  # Garante thread-safety e evita problemas com GUI em servidores (Streamlit Cloud)
import matplotlib.pyplot as plt

import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Configurações globais de cores Matplotlib
COLOR_STATUS = {
    "Aberta": "#FF4D4D",
    "Parcialmente Encerrada": "#FF9F43",
    "Homologação": "#FED330",
    "Encerrada": "#2DD4BF"
}
COLOR_COMISSAO = {
    "Comissão Atual": "#4B7BEC",
    "Nota Provisória": "#FFD166"
}

def rpm_sort_key(name):
    """Ordenação numérica para RPMs, alfabética para outros."""
    m = re.match(r'^(\d+)\s+RPM', str(name))
    if m:
        return (0, int(m.group(1)), "")
    return (1, 0, str(name))

def set_cell_background(cell, hex_color):
    """Define a cor de fundo de uma célula na tabela do Word."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Ajusta o preenchimento (padding) interno das células de uma tabela."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tc_mar.append(node)
    tc_pr.append(tc_mar)

def create_status_pie(df_sub, filename):
    """Gera o gráfico de rosquinha (donut) para Status de Avaliação."""
    plt.figure(figsize=(5.5, 4.2))
    
    ordered = ["Aberta", "Parcialmente Encerrada", "Homologação", "Encerrada"]
    counts = df_sub["Status Avaliação"].value_counts()
    
    labels = []
    sizes = []
    colors = []
    
    for status in ordered:
        qty = counts.get(status, 0)
        if qty > 0:
            labels.append(f"{status}\n({qty:,})")
            sizes.append(qty)
            colors.append(COLOR_STATUS.get(status, "#CCCCCC"))
            
    if not sizes:
        sizes = [1]
        labels = ["Sem registros"]
        colors = ["#E0E0E0"]

    # Cria o gráfico de pizza
    wedges, texts, autotexts = plt.pie(
        sizes, labels=labels, autopct='%1.1f%%', startangle=90,
        colors=colors, wedgeprops=dict(width=0.4, edgecolor='w', linewidth=2),
        pctdistance=0.75
    )
    
    # Estiliza os textos
    for t in texts:
        t.set_fontsize(9)
        t.set_fontname("Arial")
    for at in autotexts:
        at.set_fontsize(8)
        at.set_fontweight('bold')
        at.set_color('black')
        at.set_fontname("Arial")
        
    plt.title("Status das Avaliações", fontsize=11, fontweight='bold', fontname="Arial", pad=15)
    plt.tight_layout()
    plt.savefig(filename, dpi=180, bbox_inches='tight')
    plt.close()

def create_comissao_bar(df_sub, filename):
    """Gera o gráfico de barras verticais para Comissão Atual vs Nota Provisória."""
    plt.figure(figsize=(4.5, 4.2))
    
    counts = df_sub["Situação Comissão"].value_counts()
    categories = ["Comissão Atual", "Nota Provisória"]
    values = [counts.get(cat, 0) for cat in categories]
    colors = [COLOR_COMISSAO.get(cat, "#CCCCCC") for cat in categories]
    
    bars = plt.bar(categories, values, color=colors, width=0.55, edgecolor='#DDDDDD')
    
    # Remove bordas do gráfico para visual limpo
    ax = plt.gca()
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.yaxis.set_visible(False)
    
    # Rótulos nas barras
    for bar in bars:
        height = bar.get_height()
        plt.text(
            bar.get_x() + bar.get_width()/2.0,
            height + max(1, sum(values)*0.01),
            f"{int(height):,}",
            ha='center', va='bottom', fontsize=9, fontweight='bold', fontname="Arial"
        )
        
    plt.xticks(fontsize=10, fontname="Arial")
    plt.title("Situação da Comissão", fontsize=11, fontweight='bold', fontname="Arial", pad=15)
    plt.tight_layout()
    plt.savefig(filename, dpi=180, bbox_inches='tight')
    plt.close()

def create_pending_units_bar(df, udi_udg_list, filename):
    """Gera gráfico de pendências acumuladas por UDI/UDG (Em Aberto + Parcialmente Encerrada)."""
    plt.figure(figsize=(10, 4.5))
    
    # Filtra apenas registros pendentes
    df_pend = df[df["Status Avaliação"].isin(["Aberta", "Parcialmente Encerrada"])].copy()
    
    unit_totals = {}
    for unit in udi_udg_list:
        sub_df = df_pend[df_pend["Unidade RPM (Avaliado)"] == unit]
        unit_totals[unit] = len(sub_df)
        
    # Ordena da maior para a menor pendência
    sorted_units = sorted(unit_totals.items(), key=lambda x: -x[1])
    
    # Plota se houver dados
    if sorted_units and any(val > 0 for _, val in sorted_units):
        names = [x[0] for x in sorted_units]
        values = [x[1] for x in sorted_units]
        
        bars = plt.bar(names, values, color="#FF9F43", edgecolor='#E67E22', width=0.6)
        
        ax = plt.gca()
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_visible(False)
        ax.yaxis.set_visible(False)
        
        plt.xticks(rotation=45, ha='right', fontsize=9, fontname="Arial")
        
        for bar in bars:
            height = bar.get_height()
            if height > 0:
                plt.text(
                    bar.get_x() + bar.get_width()/2.0,
                    height + max(1, sum(values)*0.005),
                    f"{int(height):,}",
                    ha='center', va='bottom', fontsize=8, fontweight='bold', fontname="Arial"
                )
    else:
        plt.text(0.5, 0.5, "Nenhuma pendência pendente no Estado", ha='center', va='center', fontsize=12)
        
    plt.title("Volume de Pendências por Unidade (UDI/UDG)", fontsize=12, fontweight='bold', fontname="Arial", pad=15)
    plt.tight_layout()
    plt.savefig(filename, dpi=180, bbox_inches='tight')
    plt.close()

def generate_word_report(df_source: pd.DataFrame, report_mode: str, selected_rpms: list = None) -> bytes:
    """Processa os dados e cria o relatório Word em bytes."""
    # 1. Limpeza e Normalização dos Dados
    df = df_source.copy()
    
    # Substituir nulos ou vazios em Unidade Principal
    df["Unidade Principal (Avaliado)"] = (
        df["Unidade Principal (Avaliado)"]
        .astype(str)
        .str.strip()
        .replace({"nan": "", "-": ""})
    )
    mask_empty = df["Unidade Principal (Avaliado)"] == ""
    df.loc[mask_empty, "Unidade Principal (Avaliado)"] = df.loc[mask_empty, "Unidade RPM (Avaliado)"]
    
    # Determinar a lista de unidades principais (RPMs/UDGs)
    unique_rpms = df["Unidade RPM (Avaliado)"].dropna().unique().tolist()
    
    # Separar UDI e UDG
    udis = [u for u in unique_rpms if "RPM" in str(u)]
    udgs = [u for u in unique_rpms if "RPM" not in str(u)]
    
    udis_sorted = sorted(udis, key=rpm_sort_key)
    udgs_sorted = sorted(udgs)
    all_units_sorted = udis_sorted + udgs_sorted
    
    # Filtrar unidades se o modo for específico
    if report_mode == "especifica" and selected_rpms:
        units_to_render = [u for u in all_units_sorted if u in selected_rpms]
        # Filtrar o próprio df se for específico para refletir o escopo na Capa e Visão Geral
        df_context = df[df["Unidade RPM (Avaliado)"].isin(units_to_render)].copy()
    else:
        units_to_render = all_units_sorted
        df_context = df
        
    # Totais Globais no Contexto
    total_evals = len(df_context)
    total_enc = len(df_context[df_context["Status Avaliação"] == "Encerrada"])
    total_pend = len(df_context[df_context["Status Avaliação"].isin(["Aberta", "Parcialmente Encerrada"])])
    total_hom = len(df_context[df_context["Status Avaliação"] == "Homologação"])
    
    # 2. Inicializar Documento Word
    doc = Document()
    
    # Configurações de margens (1 polegada = 2.54 cm)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Estilos globais
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ──────────────────────────────────────────────────────────────────────────
    # CAPA DO DOCUMENTO
    # ──────────────────────────────────────────────────────────────────────────
    
    # Título Principal (Espaçamento superior artificial com parágrafos vazios)
    for _ in range(3):
        doc.add_paragraph()
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("RELATÓRIO EXECUTIVO GERENCIAL")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1F, 0x38, 0x64) # Azul Escuro PMMG
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("AADP 2026 — Avaliação de Desempenho dos Policiais Militares")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    for _ in range(4):
        doc.add_paragraph()
        
    # Painel de Resumo na Capa
    p_resumo_header = doc.add_paragraph()
    p_resumo_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_rh = p_resumo_header.add_run("RESUMO DO ESCOPO")
    run_rh.font.bold = True
    run_rh.font.size = Pt(12)
    run_rh.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    
    # Tabela estilizada na capa
    table_capa = doc.add_table(rows=5, cols=2)
    table_capa.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    
    capa_data = [
        ("Total de Avaliações", f"{total_evals:,}"),
        ("Encerradas (Concluídas)", f"{total_enc:,} ({total_enc/max(total_evals, 1)*100:.1f}%)"),
        ("Pendentes (AV1 / AV2)", f"{total_pend:,} ({total_pend/max(total_evals, 1)*100:.1f}%)"),
        ("Aguardando Homologação", f"{total_hom:,} ({total_hom/max(total_evals, 1)*100:.1f}%)"),
        ("Data de Referência", datetime.now().strftime("%d/%m/%Y")),
    ]
    
    for i, (label, val) in enumerate(capa_data):
        row = table_capa.rows[i]
        row.cells[0].paragraphs[0].add_run(label).font.bold = True
        row.cells[1].paragraphs[0].add_run(val)
        
        # Estilos das células da capa
        for cell in row.cells:
            set_cell_background(cell, "F2F5F9")
            set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
            
    for _ in range(5):
        doc.add_paragraph()
        
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_foot = p_footer.add_run(f"Polícia Militar de Minas Gerais\nDiretoria de Recursos Humanos · {datetime.now().year}")
    run_foot.font.size = Pt(9.5)
    run_foot.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    
    doc.add_page_break()

    # ──────────────────────────────────────────────────────────────────────────
    # SEÇÃO: VISÃO GERAL — ESTADO
    # ──────────────────────────────────────────────────────────────────────────
    h1 = doc.add_paragraph()
    r_h1 = h1.add_run("1. Visão Geral — Estado")
    r_h1.font.size = Pt(20)
    r_h1.font.bold = True
    r_h1.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    
    doc.add_paragraph("Esta seção apresenta o panorama geral consolidado do estado de Minas Gerais com base em todos os registros válidos do AADP 2026.")
    
    # Gerar gráficos estaduais em arquivos temporários
    fd_pie, temp_pie = tempfile.mkstemp(suffix=".png")
    os.close(fd_pie)
    fd_bar, temp_bar = tempfile.mkstemp(suffix=".png")
    os.close(fd_bar)
    fd_pend, temp_pend = tempfile.mkstemp(suffix=".png")
    os.close(fd_pend)
    
    create_status_pie(df_context, temp_pie)
    create_comissao_bar(df_context, temp_bar)
    create_pending_units_bar(df_context, units_to_render, temp_pend)
    
    # Inserir par de gráficos lado a lado em uma tabela sem bordas
    table_graphs = doc.add_table(rows=1, cols=2)
    table_graphs.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    cell_left = table_graphs.rows[0].cells[0]
    cell_right = table_graphs.rows[0].cells[1]
    
    cell_left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_left.paragraphs[0].add_run().add_picture(temp_pie, width=Inches(3.1))
    
    cell_right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_right.paragraphs[0].add_run().add_picture(temp_bar, width=Inches(2.8))
    
    # Adicionar o gráfico de barras de pendências por UDI/UDG
    p_graph_pend = doc.add_paragraph()
    p_graph_pend.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_graph_pend.add_run().add_picture(temp_pend, width=Inches(6.2))
    
    # ── TABELA 1: Subunidades 100% Concluídas ─────────────────────────────────
    doc.add_paragraph()
    h2_1 = doc.add_paragraph()
    r_h2_1 = h2_1.add_run("1.1 Subunidades 100% Encerradas")
    r_h2_1.font.size = Pt(13)
    r_h2_1.font.bold = True
    r_h2_1.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)
    
    doc.add_paragraph("Abaixo estão listadas as Subunidades que concluíram 100% de seus processos avaliativos, sem nenhuma pendência em aberto ou parcialmente encerrada.")
    
    # Agrupa por Subunidade para calcular pendências
    sub_stats = df_context.groupby(["Unidade RPM (Avaliado)", "Unidade Principal (Avaliado)"]).agg(
        total=("Status Avaliação", "count"),
        concluidas=("Status Avaliação", lambda x: sum(x == "Encerrada")),
        pendentes=("Status Avaliação", lambda x: sum(x.isin(["Aberta", "Parcialmente Encerrada"])))
    ).reset_index()
    
    sub_100 = sub_stats[sub_stats["pendentes"] == 0].sort_values("total", ascending=False)
    
    # Cria a tabela de concluídos
    table_100 = doc.add_table(rows=1, cols=3)
    table_100.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table_100.style = 'Light Shading Accent 1' # Estilo padrão limpo do Word
    
    hdr_cells = table_100.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run("UDI / UDG").font.bold = True
    hdr_cells[1].paragraphs[0].add_run("Subunidade").font.bold = True
    hdr_cells[2].paragraphs[0].add_run("Total Avaliações").font.bold = True
    
    # Cor do cabeçalho
    for cell in hdr_cells:
        set_cell_background(cell, "1F3864")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        
    if sub_100.empty:
        row_cells = table_100.add_row().cells
        row_cells[0].paragraphs[0].add_run("Nenhuma subunidade 100% encerrada no momento.")
        row_cells[1].paragraphs[0].add_run("-")
        row_cells[2].paragraphs[0].add_run("-")
    else:
        for idx, r in sub_100.iterrows():
            row_cells = table_100.add_row().cells
            row_cells[0].paragraphs[0].add_run(str(r["Unidade RPM (Avaliado)"]))
            row_cells[1].paragraphs[0].add_run(str(r["Unidade Principal (Avaliado)"]))
            row_cells[2].paragraphs[0].add_run(f"{int(r['total']):,}")
            for cell in row_cells:
                set_cell_margins(cell, top=80, bottom=80, left=150, right=150)
                
    # ── TABELA 2: Ranking de Subunidades com Pendências ───────────────────────
    doc.add_paragraph()
    h2_2 = doc.add_paragraph()
    r_h2_2 = h2_2.add_run("1.2 Ranking de Subunidades com Pendências")
    r_h2_2.font.size = Pt(13)
    r_h2_2.font.bold = True
    r_h2_2.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)
    
    doc.add_paragraph("Abaixo estão ranqueadas as Subunidades com pendências ativas (AV1 aberta ou AV2 parcial), classificadas da menor para a maior pendência.")
    
    # Subunidades com pendência > 0
    sub_pend = sub_stats[sub_stats["pendentes"] > 0].copy()
    
    # Calcula detalhadamente Em Aberto e Parcialmente Encerradas
    det_pend = df_context[df_context["Status Avaliação"].isin(["Aberta", "Parcialmente Encerrada"])].groupby(
        ["Unidade RPM (Avaliado)", "Unidade Principal (Avaliado)", "Status Avaliação"]
    ).size().unstack(fill_value=0).reset_index()
    
    # Garante que colunas existam
    for st_col in ["Aberta", "Parcialmente Encerrada"]:
        if st_col not in det_pend.columns:
            det_pend[st_col] = 0
            
    # Junta de volta com sub_stats
    sub_pend = sub_pend.merge(det_pend, on=["Unidade RPM (Avaliado)", "Unidade Principal (Avaliado)"], how="left").fillna(0)
    
    # Total de pendências do estado para calcular o percentual do passivo
    total_state_pend = df_context[df_context["Status Avaliação"].isin(["Aberta", "Parcialmente Encerrada"])].shape[0]
    
    sub_pend["Passivo PMMG %"] = (sub_pend["pendentes"] / max(total_state_pend, 1)) * 100
    
    # Ordenar da MENOR para a MAIOR pendência
    sub_pend = sub_pend.sort_values("pendentes", ascending=True)
    
    # Cria a tabela de pendentes
    table_pend = doc.add_table(rows=1, cols=6)
    table_pend.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table_pend.style = 'Light Shading Accent 1'
    
    hdr_cells = table_pend.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run("UDI / UDG").font.bold = True
    hdr_cells[1].paragraphs[0].add_run("Subunidade").font.bold = True
    hdr_cells[2].paragraphs[0].add_run("Em Aberto").font.bold = True
    hdr_cells[3].paragraphs[0].add_run("Parc. Encerrada").font.bold = True
    hdr_cells[4].paragraphs[0].add_run("Total Pendentes").font.bold = True
    hdr_cells[5].paragraphs[0].add_run("Passivo %").font.bold = True
    
    for cell in hdr_cells:
        set_cell_background(cell, "A52A2A") # Cor vermelha/marrom escuro para sinalizar alerta/pendência
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        
    if sub_pend.empty:
        row_cells = table_pend.add_row().cells
        row_cells[0].paragraphs[0].add_run("Nenhuma subunidade com pendências pendentes.")
        for i in range(1, 6):
            row_cells[i].paragraphs[0].add_run("-")
    else:
        for idx, r in sub_pend.iterrows():
            row_cells = table_pend.add_row().cells
            row_cells[0].paragraphs[0].add_run(str(r["Unidade RPM (Avaliado)"]))
            row_cells[1].paragraphs[0].add_run(str(r["Unidade Principal (Avaliado)"]))
            row_cells[2].paragraphs[0].add_run(f"{int(r['Aberta']):,}")
            row_cells[3].paragraphs[0].add_run(f"{int(r['Parcialmente Encerrada']):,}")
            row_cells[4].paragraphs[0].add_run(f"{int(r['pendentes']):,}").font.bold = True
            row_cells[5].paragraphs[0].add_run(f"{r['Passivo PMMG %']:.2f}%")
            for cell in row_cells:
                set_cell_margins(cell, top=80, bottom=80, left=150, right=150)
                
    # Limpa arquivos temporários estaduais
    for path in [temp_pie, temp_bar, temp_pend]:
        try:
            os.remove(path)
        except Exception:
            pass
            
    # ──────────────────────────────────────────────────────────────────────────
    # SEÇÕES: DETALHAMENTO DE UNIDADES (UDI / UDG)
    # ──────────────────────────────────────────────────────────────────────────
    doc.add_page_break()
    
    sec_num = 2
    for unit_name in units_to_render:
        df_unit = df[df["Unidade RPM (Avaliado)"] == unit_name].copy()
        if df_unit.empty:
            continue
            
        h_unit = doc.add_paragraph()
        r_hu = h_unit.add_run(f"{sec_num}. Detalhamento da Unidade: {unit_name}")
        r_hu.font.size = Pt(16)
        r_hu.font.bold = True
        r_hu.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        
        # Totais específicos da unidade
        u_total = len(df_unit)
        u_enc = len(df_unit[df_unit["Status Avaliação"] == "Encerrada"])
        u_pend = len(df_unit[df_unit["Status Avaliação"].isin(["Aberta", "Parcialmente Encerrada"])])
        u_hom = len(df_unit[df_unit["Status Avaliação"] == "Homologação"])
        
        doc.add_paragraph(
            f"A unidade {unit_name} possui no momento {u_total:,} avaliações processadas, das quais "
            f"{u_enc:,} estão encerradas ({u_enc/max(u_total,1)*100:.1f}%), {u_pend:,} estão pendentes "
            f"({u_pend/max(u_total,1)*100:.1f}%) e {u_hom:,} aguardam homologação."
        )
        
        # Gráficos da unidade
        fd_upie, temp_upie = tempfile.mkstemp(suffix=".png")
        os.close(fd_upie)
        fd_ubar, temp_ubar = tempfile.mkstemp(suffix=".png")
        os.close(fd_ubar)
        
        create_status_pie(df_unit, temp_upie)
        create_comissao_bar(df_unit, temp_ubar)
        
        table_ugraphs = doc.add_table(rows=1, cols=2)
        table_ugraphs.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
        cell_l = table_ugraphs.rows[0].cells[0]
        cell_r = table_ugraphs.rows[0].cells[1]
        
        cell_l.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_l.paragraphs[0].add_run().add_picture(temp_upie, width=Inches(3.0))
        
        cell_r.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_r.paragraphs[0].add_run().add_picture(temp_ubar, width=Inches(2.7))
        
        # Limpar arquivos temporários da unidade
        for path in [temp_upie, temp_ubar]:
            try:
                os.remove(path)
            except Exception:
                pass
                
        # ── SUBSEÇÃO: Unidades Subordinadas ───────────────────────────────────
        # Apenas renderizar se o modo não for "geral_rpm" (ou seja, se for completo/subordinadas ou específica)
        if report_mode != "geral_rpm":
            sub_list = df_unit["Unidade Principal (Avaliado)"].dropna().unique().tolist()
            # Ordenar subunidades por volume decrescente de avaliações
            sub_volumes = {sub: len(df_unit[df_unit["Unidade Principal (Avaliado)"] == sub]) for sub in sub_list}
            sorted_subs = sorted(sub_volumes.items(), key=lambda x: -x[1])
            
            # Se houver subunidades e não for apenas uma idêntica à própria RPM
            if sorted_subs and not (len(sorted_subs) == 1 and sorted_subs[0][0] == unit_name):
                doc.add_paragraph()
                h_sub_hdr = doc.add_paragraph()
                r_hsh = h_sub_hdr.add_run(f"{sec_num}.1 Unidades Subordinadas de {unit_name}")
                r_hsh.font.size = Pt(13)
                r_hsh.font.bold = True
                r_hsh.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)
                
                for sub_name, sub_qty in sorted_subs:
                    # Ignorar se for a própria sede se existirem outras (ou plotar todas)
                    df_sub = df_unit[df_unit["Unidade Principal (Avaliado)"] == sub_name].copy()
                    if df_sub.empty:
                        continue
                        
                    doc.add_paragraph()
                    p_sub_title = doc.add_paragraph()
                    p_sub_title.paragraph_format.keep_with_next = True
                    r_st = p_sub_title.add_run(f"■ Subunidade: {sub_name} ({sub_qty:,} avaliações)")
                    r_st.font.size = Pt(11)
                    r_st.font.bold = True
                    
                    # Gráficos da subunidade
                    fd_spie, temp_spie = tempfile.mkstemp(suffix=".png")
                    os.close(fd_spie)
                    fd_sbar, temp_sbar = tempfile.mkstemp(suffix=".png")
                    os.close(fd_sbar)
                    
                    create_status_pie(df_sub, temp_spie)
                    create_comissao_bar(df_sub, temp_sbar)
                    
                    table_sgraphs = doc.add_table(rows=1, cols=2)
                    table_sgraphs.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
                    cell_sl = table_sgraphs.rows[0].cells[0]
                    cell_sr = table_sgraphs.rows[0].cells[1]
                    
                    cell_sl.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell_sl.paragraphs[0].add_run().add_picture(temp_spie, width=Inches(2.9))
                    
                    cell_sr.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell_sr.paragraphs[0].add_run().add_picture(temp_sbar, width=Inches(2.6))
                    
                    # Limpar temporários
                    for path in [temp_spie, temp_sbar]:
                        try:
                            os.remove(path)
                        except Exception:
                            pass
                            
        doc.add_page_break()
        sec_num += 1
        
    # Salvar em stream em memória
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()
