# -*- coding: utf-8 -*-
"""
SCRIPT DE ANALISE AADP 2026 - RELATORIO PARA O ALTO COMANDO
Policia Militar de Minas Gerais - Diretoria de Recursos Humanos
Data: Julho/2026
"""

import pandas as pd
import numpy as np
import warnings
import os
from datetime import datetime

warnings.filterwarnings('ignore')

# ─────────────────────────────────────────────────────────────────────────────
# 1. CARREGAMENTO DOS DADOS
# ─────────────────────────────────────────────────────────────────────────────
print("=" * 70)
print("  ANALISE AADP 2026 - RELATORIO ALTO COMANDO PMMG")
print("=" * 70)
print("  Inicio: " + datetime.now().strftime('%d/%m/%Y %H:%M:%S'))
print("=" * 70)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

print("\n[1/5] Carregando dados do SIGEF...")
sigef = pd.read_csv(
    os.path.join(BASE_DIR, 'SIGEF.csv'),
    encoding='latin-1', sep=';', dtype=str
)
sigef.columns = [c.strip() for c in sigef.columns]
sigef.rename(columns={
    'NUMERO': 'MATRICULA_SIGEF',
    'POSTO/GRADUACAO': 'NOME_SERVIDOR_SIGEF',
    'NOME SERVIDOR': 'POSTO_GRAD_SIGEF',
    'SIT. FUNCIONAL': 'DATA_SIT_FUNC',
    'DATA SIT. FUNCIONAL': 'SIT_FUNCIONAL',
    'ATIVO/INATIVO': 'DATA_NASCIMENTO',
    'DATA NASCIMENTO': 'ATIVO_INATIVO',
}, inplace=True)
print("   -> SIGEF carregado: " + f"{len(sigef):,}" + " registros")

print("\n[2/5] Carregando Analise AADP/SIRHWEB...")
df_aadp = pd.read_excel(
    os.path.join(BASE_DIR, 'analise_avaliacoes_completa.xlsx'),
    sheet_name='Planilha', dtype=str
)
df_aadp.columns = [c.strip() for c in df_aadp.columns]
print("   -> AADP/SIRHWEB carregado: " + f"{len(df_aadp):,}" + " registros")

print("\n[3/5] Carregando Mainframe (COM_AADP_2026)...")
df_main = pd.read_excel(
    os.path.join(BASE_DIR, 'COM_AADP_2026.xlsx'),
    sheet_name='COMAADP(10)', dtype=str
)
df_main.columns = [c.strip() for c in df_main.columns]
col_motivo = [c for c in df_main.columns if 'DESCRICAO MOTIVO' in c]
if col_motivo:
    df_main.rename(columns={col_motivo[0]: 'DESCRICAO_MOTIVO'}, inplace=True)
else:
    df_main['DESCRICAO_MOTIVO'] = ''
print("   -> Mainframe carregado: " + f"{len(df_main):,}" + " registros")

# ─────────────────────────────────────────────────────────────────────────────
# 2. PRE-PROCESSAMENTO
# ─────────────────────────────────────────────────────────────────────────────
print("\n[4/5] Processando e cruzando dados...")

df_aadp['NR PM'] = pd.to_numeric(df_aadp['NR PM'], errors='coerce')
df_aadp['Qtd Avaliacoes'] = pd.to_numeric(
    df_aadp.get('Qtd Avaliações', df_aadp.get('Qtd Avaliacoes', 0)),
    errors='coerce'
).fillna(0).astype(int)
df_aadp['Nota Final'] = pd.to_numeric(
    df_aadp.get('Nota Final - Média Aritmética', df_aadp.get('Nota Final - Media Aritmetica', np.nan)),
    errors='coerce'
)
# Coluna original de Todas Avaliações Encerradas
col_enc = 'Todas Avaliações Foram Encerradas?'
if col_enc not in df_aadp.columns:
    col_enc = [c for c in df_aadp.columns if 'Encerrad' in c]
    col_enc = col_enc[0] if col_enc else None

def extrair_matricula(m):
    try:
        return int(str(m).split('-')[0].strip())
    except:
        return np.nan

df_main['MATRICULA_NUM'] = df_main['MATRICULA'].apply(extrair_matricula)
df_main['NOTA DA AADP'] = pd.to_numeric(df_main['NOTA DA AADP'], errors='coerce')

sigef['MATRICULA_SIGEF_NUM'] = pd.to_numeric(sigef['MATRICULA_SIGEF'], errors='coerce')
sigef['SIT_FUNCIONAL'] = sigef['SIT_FUNCIONAL'].fillna('').str.strip()
sigef['QUADRO'] = sigef['QUADRO'].fillna('').str.strip()

EXCLUIR_SIT = [
    'AGREG. DESERCAO', 'PRESO PREV. SEM SERV', 'CONDENADO/FAZ. SERV.',
    'AFAST DEC JUDIC C/V', 'MED.CAUT.DIV.PR.ATIV'
]
sigef_ativa = sigef[(~sigef['SIT_FUNCIONAL'].isin(EXCLUIR_SIT)) & (sigef['QUADRO'] == 'A')].copy()
print("   -> SIGEF - Militares da Ativa (Quadro A): " + f"{len(sigef_ativa):,}")

main_recon = df_main[df_main['SIT. FUNCIONAL'].str.strip() == 'PRACA QPR RECONDUZID'].copy()
main_ativos = df_main[df_main['SIT. FUNCIONAL'].str.strip() != 'PRACA QPR RECONDUZID'].copy()
print("   -> Mainframe - Reconduzidos: " + f"{len(main_recon):,}")
print("   -> Mainframe - Demais (nao reconduzidos): " + f"{len(main_ativos):,}")

recon_com_nota = main_recon[main_recon['MATRICULA_NUM'].isin(df_aadp['NR PM'].dropna())].copy()
recon_sem_nota = main_recon[~main_recon['MATRICULA_NUM'].isin(df_aadp['NR PM'].dropna())].copy()
print("   -> Reconduzidos com avaliacao no AADP/SIRHWEB: " + f"{len(recon_com_nota):,}")
print("   -> Reconduzidos SEM avaliacao no AADP/SIRHWEB: " + f"{len(recon_sem_nota):,}")

main_universo = main_ativos.copy()
main_universo['TEM_AADP'] = main_universo['MATRICULA_NUM'].isin(df_aadp['NR PM'].dropna())

# Merge
aadp_merge_cols = ['NR PM', 'Nota Final', 'Qtd Avaliacoes']
if col_enc:
    aadp_merge_cols.append(col_enc)

merged = main_universo.merge(
    df_aadp[aadp_merge_cols].rename(columns={'NR PM': 'MATRICULA_NUM'}),
    on='MATRICULA_NUM', how='left'
)

merged_com_ambas = merged[
    merged['TEM_AADP'] &
    merged['NOTA DA AADP'].notna() &
    merged['Nota Final'].notna()
].copy()

merged_com_ambas['DIFF_NOTA'] = (merged_com_ambas['Nota Final'] - merged_com_ambas['NOTA DA AADP']).round(4)
merged_com_ambas['STATUS_NOTA'] = merged_com_ambas['DIFF_NOTA'].apply(
    lambda d: 'CONVERGENTE' if abs(d) < 0.01 else 'DIVERGENTE'
)

sem_aadp = main_universo[~main_universo['TEM_AADP']].copy()
com_aadp = main_universo[main_universo['TEM_AADP']].copy()
divergentes = merged_com_ambas[merged_com_ambas['STATUS_NOTA'] == 'DIVERGENTE'].copy()
convergentes = merged_com_ambas[merged_com_ambas['STATUS_NOTA'] == 'CONVERGENTE'].copy()

# Estatisticas por situacao funcional
sit_analise = df_aadp.groupby('Sit. Funcional').agg(
    QTD_MILITARES=('NR PM', 'count'),
    MEDIA_NOTA=('Nota Final', 'mean'),
    MIN_NOTA=('Nota Final', 'min'),
    MAX_NOTA=('Nota Final', 'max'),
    COM_NOTA=('Nota Final', lambda x: x.notna().sum()),
    SEM_NOTA=('Nota Final', lambda x: x.isna().sum()),
).reset_index()
sit_analise['MEDIA_NOTA'] = sit_analise['MEDIA_NOTA'].round(2)

def faixa_nota(n):
    try:
        n = float(n)
        if n < 5.0:    return '1. Abaixo de 5,0'
        elif n < 7.0:  return '2. De 5,0 a 6,9'
        elif n < 8.0:  return '3. De 7,0 a 7,9'
        elif n < 9.0:  return '4. De 8,0 a 8,9'
        elif n < 10.0: return '5. De 9,0 a 9,9'
        else:          return '6. Nota 10,0'
    except:
        return '7. Sem Nota'

df_aadp['FAIXA_NOTA'] = df_aadp['Nota Final'].apply(faixa_nota)
dist_notas = df_aadp.groupby(['Sit. Funcional', 'FAIXA_NOTA']).size().reset_index(name='QTD')

col_rpm = 'Nome RPM'
col_enc2 = col_enc if col_enc else 'dummy_col'
if col_enc:
    rpm_analise = df_aadp.groupby(col_rpm).agg(
        QTD_MILITARES=('NR PM', 'count'),
        COM_NOTA=('Nota Final', lambda x: x.notna().sum()),
        SEM_NOTA=('Nota Final', lambda x: x.isna().sum()),
        MEDIA_NOTA=('Nota Final', 'mean'),
        ENCERRADAS=(col_enc, lambda x: (x.str.upper().str.strip() == 'SIM').sum()),
    ).reset_index()
else:
    rpm_analise = df_aadp.groupby(col_rpm).agg(
        QTD_MILITARES=('NR PM', 'count'),
        COM_NOTA=('Nota Final', lambda x: x.notna().sum()),
        SEM_NOTA=('Nota Final', lambda x: x.isna().sum()),
        MEDIA_NOTA=('Nota Final', 'mean'),
    ).reset_index()
    rpm_analise['ENCERRADAS'] = 0

rpm_analise['MEDIA_NOTA'] = rpm_analise['MEDIA_NOTA'].round(2)
rpm_analise['PCT_COM_NOTA'] = (rpm_analise['COM_NOTA'] / rpm_analise['QTD_MILITARES'] * 100).round(1)
rpm_analise['PCT_ENCERRADAS'] = (rpm_analise['ENCERRADAS'] / rpm_analise['QTD_MILITARES'] * 100).round(1)

total_mainframe   = len(df_main)
total_sigef_ativa = len(sigef_ativa)
total_universo    = len(main_universo)
total_reconduzidos= len(main_recon)
total_com_aadp    = int(main_universo['TEM_AADP'].sum())
total_sem_aadp    = len(sem_aadp)
total_comparacao  = len(merged_com_ambas)
total_convergentes= len(convergentes)
total_divergentes = len(divergentes)
pct_com_aadp      = round(total_com_aadp / total_universo * 100, 1) if total_universo else 0
pct_sem_aadp      = round(total_sem_aadp / total_universo * 100, 1) if total_universo else 0
pct_convergentes  = round(total_convergentes / total_comparacao * 100, 1) if total_comparacao else 0
pct_divergentes   = round(total_divergentes / total_comparacao * 100, 1) if total_comparacao else 0

print("\n" + "=" * 60)
print("  RESUMO EXECUTIVO")
print("=" * 60)
print(f"  Total Mainframe (Universo base):   {total_mainframe:>8,}")
print(f"  SIGEF - Militares Ativa:           {total_sigef_ativa:>8,}")
print(f"  Reconduzidos QPR (excluidos):      {total_reconduzidos:>8,}")
print(f"  Universo para analise:             {total_universo:>8,}")
print(f"  Com avaliacao AADP/SIRHWEB:        {total_com_aadp:>8,} ({pct_com_aadp}%)")
print(f"  Sem avaliacao AADP/SIRHWEB:        {total_sem_aadp:>8,} ({pct_sem_aadp}%)")
print(f"  Pares com notas nos 2 sistemas:    {total_comparacao:>8,}")
print(f"  Notas CONVERGENTES:                {total_convergentes:>8,} ({pct_convergentes}%)")
print(f"  Notas DIVERGENTES:                 {total_divergentes:>8,} ({pct_divergentes}%)")
print("=" * 60)

# ─────────────────────────────────────────────────────────────────────────────
# 3. GERACAO DO EXCEL
# ─────────────────────────────────────────────────────────────────────────────
print("\n[5/5] Gerando relatorios...")
EXCEL_OUT = os.path.join(BASE_DIR, 'RELATORIO_ALTO_COMANDO_AADP2026.xlsx')

from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter

COR_AZUL_ESC = "1B2A4A"
COR_AZUL_MED = "2E4A7A"
COR_AZUL_CLR = "4A7AB8"
COR_DOURADO  = "C8A84B"
COR_BRANCO   = "FFFFFF"
COR_CZ_CLR   = "F2F4F8"
COR_VERDE    = "1A6B3A"
COR_VERM     = "8B1A1A"
COR_LRNJ     = "C85A00"
COR_VD_CLR   = "D4EDDA"
COR_VM_CLR   = "F8D7DA"
COR_AM_CLR   = "FFF3CD"

def fx(cor):  return PatternFill("solid", fgColor=cor)
def fn(bold=False, cor="000000", sz=11, italic=False):
    return Font(bold=bold, color=cor, size=sz, italic=italic, name="Calibri")
def al(h="center", v="center", wrap=True):
    return Alignment(horizontal=h, vertical=v, wrap_text=wrap)
def bd():
    s = Side(style="thin")
    return Border(left=s, right=s, top=s, bottom=s)

wb = Workbook()

def cabecalho(ws, titulo, sub=""):
    ws.sheet_view.showGridLines = False
    ws.column_dimensions['A'].width = 3
    ws.merge_cells('B1:D1')
    ws['B1'] = "POLICIA MILITAR DE MINAS GERAIS"
    ws['B1'].font = fn(True, COR_DOURADO, 14)
    ws['B1'].fill = fx(COR_AZUL_ESC)
    ws['B1'].alignment = al()
    ws['B1'].border = bd()
    ws.row_dimensions[1].height = 28
    ws.merge_cells('B2:D2')
    ws['B2'] = "DIRETORIA DE RECURSOS HUMANOS | DRH"
    ws['B2'].font = fn(True, COR_BRANCO, 11)
    ws['B2'].fill = fx(COR_AZUL_MED)
    ws['B2'].alignment = al()
    ws['B2'].border = bd()
    ws.row_dimensions[2].height = 20
    ws.merge_cells('B3:D3')
    ws['B3'] = titulo
    ws['B3'].font = fn(True, COR_BRANCO, 12)
    ws['B3'].fill = fx(COR_AZUL_CLR)
    ws['B3'].alignment = al()
    ws['B3'].border = bd()
    ws.row_dimensions[3].height = 22
    if sub:
        ws.merge_cells('B4:D4')
        ws['B4'] = sub
        ws['B4'].font = fn(italic=True, cor=COR_AZUL_ESC, sz=10)
        ws['B4'].alignment = al()
        ws.row_dimensions[4].height = 16

# ── ABA 1: RESUMO EXECUTIVO
ws1 = wb.active
ws1.title = "1. RESUMO EXECUTIVO"
ws1.column_dimensions['B'].width = 44
ws1.column_dimensions['C'].width = 18
ws1.column_dimensions['D'].width = 14
ws1.column_dimensions['E'].width = 3
cabecalho(ws1, "ANALISE CENARIO AVALIACOES AADP 2026",
          "Referencia: " + datetime.now().strftime('%d/%m/%Y') + " | Confidencial - Alto Comando")

kpis = [
    ("TOTAL MAINFRAME (Base)", total_mainframe, COR_AZUL_MED, COR_BRANCO),
    ("SIGEF - Militares da Ativa (Quadro A)", total_sigef_ativa, COR_AZUL_ESC, COR_BRANCO),
    ("Reconduzidos QPR (excluidos da analise)", total_reconduzidos, COR_LRNJ, COR_BRANCO),
    ("Universo para Analise AADP", total_universo, COR_AZUL_CLR, COR_BRANCO),
    ("Com Avaliacao AADP/SIRHWEB", total_com_aadp, COR_VERDE, COR_BRANCO),
    ("Sem Avaliacao AADP/SIRHWEB", total_sem_aadp, COR_VERM, COR_BRANCO),
    ("Notas Convergentes (" + str(pct_convergentes) + "%)", total_convergentes, COR_VERDE, COR_BRANCO),
    ("Notas Divergentes (" + str(pct_divergentes) + "%)", total_divergentes, COR_VERM, COR_BRANCO),
]

row = 6
for (label, valor, cf, ct) in kpis:
    ws1.merge_cells('B' + str(row) + ':C' + str(row))
    ws1['B' + str(row)] = label
    ws1['B' + str(row)].font = fn(False, ct, 11)
    ws1['B' + str(row)].fill = fx(cf)
    ws1['B' + str(row)].alignment = al(h="left")
    ws1['B' + str(row)].border = bd()
    ws1['D' + str(row)] = valor
    ws1['D' + str(row)].font = fn(True, ct, 13)
    ws1['D' + str(row)].fill = fx(cf)
    ws1['D' + str(row)].alignment = al()
    ws1['D' + str(row)].border = bd()
    ws1.row_dimensions[row].height = 26
    row += 1

row += 1
ws1.merge_cells('B' + str(row) + ':D' + str(row))
ws1['B' + str(row)] = ">> SITUACAO FUNCIONAL DOS MILITARES NO MAINFRAME (excl. reconduzidos)"
ws1['B' + str(row)].font = fn(True, COR_BRANCO, 11)
ws1['B' + str(row)].fill = fx(COR_AZUL_ESC)
ws1['B' + str(row)].alignment = al(h="left")
ws1['B' + str(row)].border = bd()
ws1.row_dimensions[row].height = 22
row += 1

sit_counts = main_universo['SIT. FUNCIONAL'].value_counts().reset_index()
sit_counts.columns = ['SIT. FUNCIONAL', 'QTD']
sit_counts['PCT'] = (sit_counts['QTD'] / sit_counts['QTD'].sum() * 100).round(1)

for col, lbl in [('B', 'SITUACAO FUNCIONAL'), ('C', 'QTD. MILITARES'), ('D', '% DO TOTAL')]:
    ws1[col + str(row)] = lbl
    ws1[col + str(row)].font = fn(True, COR_BRANCO, 10)
    ws1[col + str(row)].fill = fx(COR_AZUL_CLR)
    ws1[col + str(row)].alignment = al()
    ws1[col + str(row)].border = bd()
ws1.row_dimensions[row].height = 18
row += 1

for i, (_, r) in enumerate(sit_counts.iterrows()):
    c = COR_CZ_CLR if i % 2 == 0 else COR_BRANCO
    ws1['B' + str(row)] = str(r['SIT. FUNCIONAL']).strip()
    ws1['C' + str(row)] = int(r['QTD'])
    ws1['D' + str(row)] = str(r['PCT']) + "%"
    for col in ['B','C','D']:
        ws1[col + str(row)].font = fn(sz=10)
        ws1[col + str(row)].fill = fx(c)
        ws1[col + str(row)].alignment = al(h="left" if col == 'B' else "center")
        ws1[col + str(row)].border = bd()
    ws1.row_dimensions[row].height = 17
    row += 1

# ── ABA 2: SEM AVALIACAO AADP
ws2 = wb.create_sheet("2. SEM AVALIACAO AADP")
cabecalho(ws2, "MILITARES SEM AVALIACAO NO AADP/SIRHWEB",
          "Total: " + f"{total_sem_aadp:,}" + " militares | Excluidos reconduzidos")
for col, w in zip(list('BCDEFGHIJ'), [14, 28, 36, 34, 34, 22, 8, 22, 14]):
    ws2.column_dimensions[col].width = w

sem_exp_cols = ['MATRICULA','POSTO/GRADUACAO','NOME SERVIDOR','NOME UNIDADE PRINCIPAL',
                'NOME UNIDADE','NOME MUNICIPIO','QUADRO','SIT. FUNCIONAL','NOTA DA AADP']
sem_exp_hdrs = ['MATRICULA','POSTO/GRAD','NOME SERVIDOR','UNIDADE PRINCIPAL',
                'UNIDADE','MUNICIPIO','QUADRO','SIT. FUNCIONAL','NOTA MAINFRAME']

row = 6
for cl, h in zip(list('BCDEFGHIJ'), sem_exp_hdrs):
    ws2[cl + str(row)] = h
    ws2[cl + str(row)].font = fn(True, COR_BRANCO, 10)
    ws2[cl + str(row)].fill = fx(COR_AZUL_MED)
    ws2[cl + str(row)].alignment = al()
    ws2[cl + str(row)].border = bd()
ws2.row_dimensions[row].height = 20
row += 1

sem_aadp_exp = sem_aadp[[c for c in sem_exp_cols if c in sem_aadp.columns]]
for i, (_, r) in enumerate(sem_aadp_exp.iterrows()):
    c = COR_CZ_CLR if i % 2 == 0 else COR_BRANCO
    for cl, col in zip(list('BCDEFGHIJ'), sem_exp_cols):
        if col in r.index:
            ws2[cl + str(row)] = str(r[col]).strip() if pd.notna(r[col]) else ''
        else:
            ws2[cl + str(row)] = ''
        ws2[cl + str(row)].font = fn(sz=9)
        ws2[cl + str(row)].fill = fx(c)
        ws2[cl + str(row)].alignment = al(h="left", wrap=False)
        ws2[cl + str(row)].border = bd()
    ws2.row_dimensions[row].height = 15
    row += 1

# ── ABA 3: CONVERGENCIA DE NOTAS
ws3 = wb.create_sheet("3. CONVERGENCIA DE NOTAS")
cabecalho(ws3, "COMPARACAO NOTAS AADP/SIRHWEB x MAINFRAME",
          "Convergentes: " + f"{total_convergentes:,}" + " | Divergentes: " + f"{total_divergentes:,}")
for col, w in zip(list('BCDEFGHI'), [14, 36, 34, 22, 16, 18, 12, 15]):
    ws3.column_dimensions[col].width = w

c3_hdrs = ['MATRICULA','NOME SERVIDOR','UNIDADE PRINCIPAL','SIT. FUNCIONAL',
           'NOTA MAINFRAME','NOTA AADP/SIRHWEB','DIFERENCA','STATUS']
row = 6
for cl, h in zip(list('BCDEFGHI'), c3_hdrs):
    ws3[cl + str(row)] = h
    ws3[cl + str(row)].font = fn(True, COR_BRANCO, 10)
    ws3[cl + str(row)].fill = fx(COR_AZUL_MED)
    ws3[cl + str(row)].alignment = al()
    ws3[cl + str(row)].border = bd()
ws3.row_dimensions[row].height = 20
row += 1

for i, (_, r) in enumerate(merged_com_ambas.iterrows()):
    st = str(r.get('STATUS_NOTA', '')).strip()
    c = COR_VM_CLR if st == 'DIVERGENTE' else COR_VD_CLR
    cs = COR_VERM if st == 'DIVERGENTE' else COR_VERDE
    vals = [
        str(r.get('MATRICULA', '')).strip(),
        str(r.get('NOME SERVIDOR', '')).strip(),
        str(r.get('NOME UNIDADE PRINCIPAL', '')).strip(),
        str(r.get('SIT. FUNCIONAL', '')).strip(),
        r.get('NOTA DA AADP', ''),
        r.get('Nota Final', ''),
        r.get('DIFF_NOTA', ''),
        st,
    ]
    for cl, v in zip(list('BCDEFGHI'), vals):
        ws3[cl + str(row)] = v
        ws3[cl + str(row)].font = fn(sz=9, cor=cs if cl == 'I' else "000000", bold=(cl == 'I'))
        ws3[cl + str(row)].fill = fx(c)
        ws3[cl + str(row)].alignment = al(h="center" if cl in list('FGHI') else "left", wrap=False)
        ws3[cl + str(row)].border = bd()
    ws3.row_dimensions[row].height = 15
    row += 1

# ── ABA 4: RECONDUZIDOS QPR
ws4 = wb.create_sheet("4. RECONDUZIDOS QPR")
cabecalho(ws4, "RECONDUZIDOS QPR - ANALISE DETALHADA",
          "Com aval AADP: " + f"{len(recon_com_nota):,}" + " | Sem aval AADP: " + f"{len(recon_sem_nota):,}")
for col, w in zip(list('BCDEFGHI'), [14, 36, 34, 22, 16, 18, 12, 16]):
    ws4.column_dimensions[col].width = w

main_recon2 = main_recon.merge(
    df_aadp[['NR PM', 'Nota Final', 'Qtd Avaliacoes']].rename(columns={'NR PM': 'MATRICULA_NUM'}),
    on='MATRICULA_NUM', how='left'
)
main_recon2['STATUS_AADP'] = main_recon2['Nota Final'].apply(
    lambda x: 'COM AVALIACAO' if pd.notna(x) else 'SEM AVALIACAO'
)

r4_hdrs = ['MATRICULA','NOME SERVIDOR','UNIDADE PRINCIPAL','SIT. FUNCIONAL',
           'NOTA MAINFRAME','NOTA AADP/SIRHWEB','QTD AVAL.','STATUS AVAL.']
r4_cols = ['MATRICULA','NOME SERVIDOR','NOME UNIDADE PRINCIPAL','SIT. FUNCIONAL',
           'NOTA DA AADP','Nota Final','Qtd Avaliacoes','STATUS_AADP']

row = 6
for cl, h in zip(list('BCDEFGHI'), r4_hdrs):
    ws4[cl + str(row)] = h
    ws4[cl + str(row)].font = fn(True, COR_BRANCO, 10)
    ws4[cl + str(row)].fill = fx(COR_LRNJ)
    ws4[cl + str(row)].alignment = al()
    ws4[cl + str(row)].border = bd()
ws4.row_dimensions[row].height = 20
row += 1

for i, (_, r) in enumerate(main_recon2.iterrows()):
    st = str(r.get('STATUS_AADP', '')).strip()
    c = COR_VD_CLR if st == 'COM AVALIACAO' else COR_AM_CLR
    for cl, col in zip(list('BCDEFGHI'), r4_cols):
        v = r.get(col, '')
        ws4[cl + str(row)] = v if pd.notna(v) else ''
        ws4[cl + str(row)].font = fn(sz=9)
        ws4[cl + str(row)].fill = fx(c)
        ws4[cl + str(row)].alignment = al(h="left" if cl in ['B','C','D'] else "center", wrap=False)
        ws4[cl + str(row)].border = bd()
    ws4.row_dimensions[row].height = 15
    row += 1

# ── ABA 5: ANALISE POR RPM
ws5 = wb.create_sheet("5. ANALISE POR RPM")
cabecalho(ws5, "ANALISE POR REGIAO DE POLICIA MILITAR (RPM)",
          "Cobertura de avaliacoes e notas medias por RPM")
for col, w in zip(list('BCDEFGHI'), [40, 16, 12, 12, 14, 14, 14, 14]):
    ws5.column_dimensions[col].width = w

r5_hdrs = ['RPM','TOTAL MILITARES','COM NOTA','SEM NOTA','% COM NOTA','MEDIA NOTA','ENCERRADAS','% ENCERRADAS']
row = 6
for cl, h in zip(list('BCDEFGHI'), r5_hdrs):
    ws5[cl + str(row)] = h
    ws5[cl + str(row)].font = fn(True, COR_BRANCO, 10)
    ws5[cl + str(row)].fill = fx(COR_AZUL_MED)
    ws5[cl + str(row)].alignment = al()
    ws5[cl + str(row)].border = bd()
ws5.row_dimensions[row].height = 20
row += 1

rpm_sorted = rpm_analise.sort_values('QTD_MILITARES', ascending=False)
for i, (_, r) in enumerate(rpm_sorted.iterrows()):
    c = COR_CZ_CLR if i % 2 == 0 else COR_BRANCO
    pct = r.get('PCT_COM_NOTA', 0)
    cpct = COR_VD_CLR if pct >= 80 else (COR_AM_CLR if pct >= 50 else COR_VM_CLR)
    vals = [
        str(r['Nome RPM']).strip(),
        int(r['QTD_MILITARES']),
        int(r['COM_NOTA']),
        int(r['SEM_NOTA']),
        str(r['PCT_COM_NOTA']) + "%",
        round(float(r['MEDIA_NOTA']), 2) if pd.notna(r['MEDIA_NOTA']) else '-',
        int(r['ENCERRADAS']),
        str(r['PCT_ENCERRADAS']) + "%",
    ]
    for cl, v in zip(list('BCDEFGHI'), vals):
        ws5[cl + str(row)] = v
        ws5[cl + str(row)].font = fn(sz=9)
        ws5[cl + str(row)].fill = fx(cpct if cl == 'F' else c)
        ws5[cl + str(row)].alignment = al(h="left" if cl == 'B' else "center", wrap=False)
        ws5[cl + str(row)].border = bd()
    ws5.row_dimensions[row].height = 15
    row += 1

# ── ABA 6: SITUACAO FUNCIONAL
ws6 = wb.create_sheet("6. SITUACAO FUNCIONAL")
cabecalho(ws6, "ANALISE POR SITUACAO FUNCIONAL",
          "Estatisticas de avaliacao e notas por situacao funcional")
for col, w in zip(list('BCDEFGH'), [30, 10, 12, 12, 14, 14, 14]):
    ws6.column_dimensions[col].width = w

r6_hdrs = ['SITUACAO FUNCIONAL','TOTAL','COM NOTA','SEM NOTA','MEDIA NOTA','MENOR NOTA','MAIOR NOTA']
row = 6
for cl, h in zip(list('BCDEFGH'), r6_hdrs):
    ws6[cl + str(row)] = h
    ws6[cl + str(row)].font = fn(True, COR_BRANCO, 10)
    ws6[cl + str(row)].fill = fx(COR_AZUL_MED)
    ws6[cl + str(row)].alignment = al()
    ws6[cl + str(row)].border = bd()
ws6.row_dimensions[row].height = 20
row += 1

for i, (_, r) in enumerate(sit_analise.iterrows()):
    c = COR_CZ_CLR if i % 2 == 0 else COR_BRANCO
    vals = [
        str(r['Sit. Funcional']).strip(),
        int(r['QTD_MILITARES']),
        int(r['COM_NOTA']),
        int(r['SEM_NOTA']),
        round(float(r['MEDIA_NOTA']), 2) if pd.notna(r['MEDIA_NOTA']) else '-',
        round(float(r['MIN_NOTA']), 2) if pd.notna(r['MIN_NOTA']) else '-',
        round(float(r['MAX_NOTA']), 2) if pd.notna(r['MAX_NOTA']) else '-',
    ]
    for cl, v in zip(list('BCDEFGH'), vals):
        ws6[cl + str(row)] = v
        ws6[cl + str(row)].font = fn(sz=10)
        ws6[cl + str(row)].fill = fx(c)
        ws6[cl + str(row)].alignment = al(h="left" if cl == 'B' else "center", wrap=False)
        ws6[cl + str(row)].border = bd()
    ws6.row_dimensions[row].height = 17
    row += 1

# ── ABA 7: DIVERGENCIAS DETALHADAS
ws7 = wb.create_sheet("7. DIVERGENCIAS DETALHADAS")
cabecalho(ws7, "DIVERGENCIAS DE NOTAS - ANALISE DETALHADA",
          "Total de divergencias: " + f"{total_divergentes:,}" + " registros")
for col, w in zip(list('BCDEFGHIJ'), [14, 36, 34, 22, 16, 18, 12, 15, 30]):
    ws7.column_dimensions[col].width = w

r7_hdrs = ['MATRICULA','NOME SERVIDOR','UNIDADE PRINCIPAL','SIT. FUNCIONAL',
           'NOTA MAINFRAME','NOTA AADP','DIFERENCA','STATUS','DESC. MOTIVO']
row = 6
for cl, h in zip(list('BCDEFGHIJ'), r7_hdrs):
    ws7[cl + str(row)] = h
    ws7[cl + str(row)].font = fn(True, COR_BRANCO, 10)
    ws7[cl + str(row)].fill = fx(COR_VERM)
    ws7[cl + str(row)].alignment = al()
    ws7[cl + str(row)].border = bd()
ws7.row_dimensions[row].height = 20
row += 1

for i, (_, r) in enumerate(divergentes.iterrows()):
    c = COR_VM_CLR if i % 2 == 0 else COR_BRANCO
    vals = [
        str(r.get('MATRICULA', '')).strip(),
        str(r.get('NOME SERVIDOR', '')).strip(),
        str(r.get('NOME UNIDADE PRINCIPAL', '')).strip(),
        str(r.get('SIT. FUNCIONAL', '')).strip(),
        r.get('NOTA DA AADP', ''),
        r.get('Nota Final', ''),
        r.get('DIFF_NOTA', ''),
        str(r.get('STATUS_NOTA', '')).strip(),
        str(r.get('DESCRICAO_MOTIVO', '')).strip()[:60],
    ]
    for cl, v in zip(list('BCDEFGHIJ'), vals):
        ws7[cl + str(row)] = v
        ws7[cl + str(row)].font = fn(sz=9)
        ws7[cl + str(row)].fill = fx(c)
        ws7[cl + str(row)].alignment = al(h="left" if cl in ['B','C','D','J'] else "center", wrap=False)
        ws7[cl + str(row)].border = bd()
    ws7.row_dimensions[row].height = 15
    row += 1

# ── ABA 8: DISTRIBUICAO DE NOTAS
ws8 = wb.create_sheet("8. DISTRIBUICAO DE NOTAS")
cabecalho(ws8, "DISTRIBUICAO DE NOTAS POR FAIXA",
          "Situacao das avaliacoes AADP/SIRHWEB por faixa de nota")

dist_pivot = dist_notas.pivot(
    index='Sit. Funcional', columns='FAIXA_NOTA', values='QTD'
).fillna(0).astype(int).reset_index()
dist_cols = list(dist_pivot.columns)
dist_widths = [28] + [14] * (len(dist_cols)-1)
for i, w in enumerate(dist_widths):
    ws8.column_dimensions[get_column_letter(i+2)].width = w

row = 6
for i, hdr in enumerate(dist_cols):
    cl = get_column_letter(i+2)
    ws8[cl + str(row)] = str(hdr).replace('Sit. Funcional', 'SITUACAO FUNCIONAL')
    ws8[cl + str(row)].font = fn(True, COR_BRANCO, 9)
    ws8[cl + str(row)].fill = fx(COR_AZUL_MED)
    ws8[cl + str(row)].alignment = al()
    ws8[cl + str(row)].border = bd()
ws8.row_dimensions[row].height = 20
row += 1

for i, (_, r) in enumerate(dist_pivot.iterrows()):
    c = COR_CZ_CLR if i % 2 == 0 else COR_BRANCO
    for j, col in enumerate(dist_cols):
        cl = get_column_letter(j+2)
        ws8[cl + str(row)] = r[col]
        ws8[cl + str(row)].font = fn(sz=9)
        ws8[cl + str(row)].fill = fx(c)
        ws8[cl + str(row)].alignment = al(h="left" if j == 0 else "center", wrap=False)
        ws8[cl + str(row)].border = bd()
    ws8.row_dimensions[row].height = 16
    row += 1

wb.save(EXCEL_OUT)
print("   Excel salvo: " + EXCEL_OUT)

# ─────────────────────────────────────────────────────────────────────────────
# 4. GERACAO DO WORD
# ─────────────────────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORD_OUT = os.path.join(BASE_DIR, 'RELATORIO_ALTO_COMANDO_AADP2026.docx')
doc = Document()

section = doc.sections[0]
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(1.5)

def rgb(hex_str):
    h = hex_str.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.upper())
    tcPr.append(shd)

def cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for edge in ('left','top','right','bottom'):
        tag = OxmlElement('w:' + edge)
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), 'AAAAAA')
        tcB.append(tag)
    tcPr.append(tcB)

def titulo(doc, texto, nivel=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = True
    if nivel == 1:
        run.font.size = Pt(16); run.font.color.rgb = rgb(COR_AZUL_ESC)
    elif nivel == 2:
        run.font.size = Pt(13); run.font.color.rgb = rgb(COR_AZUL_MED)
    else:
        run.font.size = Pt(11); run.font.color.rgb = rgb(COR_AZUL_CLR)
    run.font.name = 'Calibri'

def paragrafo(doc, texto, bold=False, sz=10, cor="000000", align="left"):
    p = doc.add_paragraph()
    aligns = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER,
              'justify': WD_ALIGN_PARAGRAPH.JUSTIFY}
    p.alignment = aligns.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(texto)
    run.bold = bold; run.font.size = Pt(sz)
    run.font.color.rgb = rgb(cor); run.font.name = 'Calibri'
    return p

def tabela(doc, dados, headers, cor_hdr=COR_AZUL_MED, widths=None):
    t = doc.add_table(rows=1+len(dados), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    hc = t.rows[0].cells
    for i, h in enumerate(headers):
        hc[i].text = str(h)
        cell_bg(hc[i], cor_hdr)
        cell_border(hc[i])
        para = hc[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = rgb(COR_BRANCO); run.font.name = 'Calibri'
    for ri, linha in enumerate(dados):
        rc = t.rows[ri+1].cells
        cc = COR_CZ_CLR if ri % 2 == 0 else COR_BRANCO
        for ci, val in enumerate(linha):
            rc[ci].text = str(val) if pd.notna(val) else '-'
            cell_bg(rc[ci], cc); cell_border(rc[ci])
            para = rc[ci].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.font.size = Pt(8); run.font.name = 'Calibri'
    if widths:
        for i, col in enumerate(t.columns):
            if i < len(widths):
                for cell in col.cells:
                    cell.width = Cm(widths[i])
    return t

# CAPA
doc.add_paragraph(); doc.add_paragraph()
titulo(doc, "POLICIA MILITAR DE MINAS GERAIS", 1)
titulo(doc, "DIRETORIA DE RECURSOS HUMANOS - DRH", 2)
doc.add_paragraph()
titulo(doc, "RELATORIO DE ANALISE", 2)
titulo(doc, "AVALIACOES DE DESEMPENHO AADP 2026", 1)
doc.add_paragraph()
paragrafo(doc, "Data de referencia: " + datetime.now().strftime('%d/%m/%Y'), align="center", sz=11)
paragrafo(doc, "CONFIDENCIAL - DESTINADO AO ALTO COMANDO", bold=True, align="center", sz=12, cor=COR_VERM)
doc.add_paragraph()
paragrafo(doc, "_" * 90, align="center", cor="AAAAAA")
doc.add_paragraph()
paragrafo(doc, "BASES DE DADOS UTILIZADAS", bold=True, sz=12, cor=COR_AZUL_ESC)
bases = [
    ("SIGEF", f"{len(sigef):,} registros", "Sistema de Gestao de Efetivo - militares da ativa (Quadro A)"),
    ("AADP/SIRHWEB", f"{len(df_aadp):,} registros", "Sistema de Avaliacoes de Desempenho Policial - notas extraidas"),
    ("MAINFRAME (COM_AADP_2026)", f"{len(df_main):,} registros", "Notas homologadas registradas no Mainframe da PMMG"),
]
for nome, qtd, desc in bases:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(nome + ": "); r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(10)
    r2 = p.add_run(qtd + " - " + desc); r2.font.name = 'Calibri'; r2.font.size = Pt(10)

doc.add_page_break()

# SECAO 1: SUMARIO EXECUTIVO
titulo(doc, "1. SUMARIO EXECUTIVO", 2)
paragrafo(doc, "_" * 90, cor="AAAAAA"); doc.add_paragraph()
txt = ("O presente relatorio apresenta a analise do cenario das avaliacoes de desempenho (AADP) "
       "do ano de 2026 para os militares da ativa da Policia Militar de Minas Gerais. "
       "A analise cruzou tres bases de dados: SIGEF (efetivo da ativa), AADP/SIRHWEB (avaliacoes "
       "e notas) e Mainframe (notas homologadas).\n\n"
       "Do universo de " + f"{total_mainframe:,}" + " militares registrados no Mainframe, foram excluidos os "
       + f"{total_reconduzidos:,}" + " Pracas QPR Reconduzidos, resultando em um universo de analise de "
       + f"{total_universo:,}" + " militares. Desse total, " + f"{total_com_aadp:,}" + " ("
       + str(pct_com_aadp) + "%) possuem avaliacao registrada no AADP/SIRHWEB, enquanto "
       + f"{total_sem_aadp:,}" + " (" + str(pct_sem_aadp) + "%) ainda nao possuem avaliacao nesse sistema.\n\n"
       "Na comparacao das notas entre os dois sistemas, foram identificados " + f"{total_comparacao:,}"
       + " militares com notas em ambas as bases. Desses, " + f"{total_convergentes:,}" + " ("
       + str(pct_convergentes) + "%) apresentaram notas convergentes (diferenca < 0,01 ponto), e "
       + f"{total_divergentes:,}" + " (" + str(pct_divergentes)
       + "%) registraram divergencia entre os sistemas, demandando atencao e regularizacao.")
paragrafo(doc, txt, sz=10, align="justify"); doc.add_paragraph()

kpi_dados = [
    ["Total Mainframe (base)", f"{total_mainframe:,}"],
    ["SIGEF - Militares Ativa (Quadro A)", f"{total_sigef_ativa:,}"],
    ["Reconduzidos QPR (excluidos da analise)", f"{total_reconduzidos:,}"],
    ["Universo para Analise AADP", f"{total_universo:,}"],
    ["Com avaliacao AADP/SIRHWEB", f"{total_com_aadp:,} ({pct_com_aadp}%)"],
    ["Sem avaliacao AADP/SIRHWEB", f"{total_sem_aadp:,} ({pct_sem_aadp}%)"],
    ["Notas comparadas (ambos sistemas)", f"{total_comparacao:,}"],
    ["Notas Convergentes", f"{total_convergentes:,} ({pct_convergentes}%)"],
    ["Notas Divergentes (requer acao)", f"{total_divergentes:,} ({pct_divergentes}%)"],
]
tabela(doc, kpi_dados, ["INDICADOR", "VALOR"], widths=[11, 5])
doc.add_paragraph()

# SECAO 2: SEM AVALIACAO
titulo(doc, "2. MILITARES SEM AVALIACAO NO SISTEMA AADP/SIRHWEB", 2)
paragrafo(doc, "_" * 90, cor="AAAAAA"); doc.add_paragraph()
txt2 = ("Foram identificados " + f"{total_sem_aadp:,}" + " militares (" + str(pct_sem_aadp) + "% do universo) "
        "que constam no Mainframe mas nao possuem avaliacao registrada no sistema AADP/SIRHWEB. "
        "Essa situacao exige atencao prioritaria, pois compromete a integridade do processo avaliativo "
        "e pode impactar promocoes, remuneracao e registros funcionais. "
        "A relacao completa desses militares esta disponivel na aba '2. SEM AVALIACAO AADP' do Excel anexo.")
paragrafo(doc, txt2, sz=10, align="justify"); doc.add_paragraph()

rpm_sem = rpm_analise.sort_values('SEM_NOTA', ascending=False).head(15)
paragrafo(doc, "RPMs com maior quantidade de militares sem avaliacao:", bold=True, sz=10)
rpm_tab_d = [[str(r['Nome RPM']).strip(), f"{int(r['QTD_MILITARES']):,}",
              f"{int(r['COM_NOTA']):,}", f"{int(r['SEM_NOTA']):,}", str(r['PCT_COM_NOTA']) + "%"]
             for _, r in rpm_sem.iterrows()]
tabela(doc, rpm_tab_d, ["RPM", "TOTAL", "COM AVAL.", "SEM AVAL.", "% COBERT."],
       widths=[7.5, 2.5, 3, 3, 3])

doc.add_page_break()

# SECAO 3: RECONDUZIDOS
titulo(doc, "3. PRACAS QPR RECONDUZIDOS", 2)
paragrafo(doc, "_" * 90, cor="AAAAAA"); doc.add_paragraph()
txt3 = ("O universo de Pracas QPR Reconduzidos compreende " + f"{total_reconduzidos:,}" + " militares. "
        "Desse total, " + f"{len(recon_com_nota):,}" + " ja possuem avaliacao registrada no sistema "
        "AADP/SIRHWEB, estando em processo avaliativo regular. Os demais "
        + f"{len(recon_sem_nota):,}" + " reconduzidos nao possuem avaliacao no AADP/SIRHWEB. "
        "Os reconduzidos que ja receberam nota ou estao sendo avaliados no AADP/SIRHWEB foram "
        "EXCLUIDOS do universo principal de analise, evitando duplicidade na contagem.")
paragrafo(doc, txt3, sz=10, align="justify"); doc.add_paragraph()
tabela(doc, [
    ["Total Reconduzidos QPR", f"{total_reconduzidos:,}"],
    ["Com avaliacao no AADP/SIRHWEB", f"{len(recon_com_nota):,}"],
    ["Sem avaliacao no AADP/SIRHWEB", f"{len(recon_sem_nota):,}"],
], ["SITUACAO", "QUANTIDADE"], widths=[10, 5])
doc.add_paragraph()

# SECAO 4: CONVERGENCIA DE NOTAS
titulo(doc, "4. CONVERGENCIA DE NOTAS - AADP/SIRHWEB x MAINFRAME", 2)
paragrafo(doc, "_" * 90, cor="AAAAAA"); doc.add_paragraph()
txt4 = ("A verificacao de convergencia foi realizada cruzando as notas do AADP/SIRHWEB "
        "(Media Aritmetica Final) com as notas homologadas no Mainframe (NOTA DA AADP). "
        "Foram identificados " + f"{total_comparacao:,}" + " militares com notas em ambos os sistemas.\n\n"
        "Criterio de convergencia: diferenca absoluta inferior a 0,01 ponto.\n"
        "  * CONVERGENTES: " + f"{total_convergentes:,}" + " militares (" + str(pct_convergentes)
        + "%) - notas identicas ou diferenca desprezivel.\n"
        "  * DIVERGENTES: " + f"{total_divergentes:,}" + " militares (" + str(pct_divergentes)
        + "%) - diferenca requer revisao e regularizacao urgente.\n\n"
        "Os casos divergentes podem indicar: lancamentos incorretos, avaliacoes nao homologadas, "
        "recursos nao processados ou erros de digitacao. Recomenda-se auditoria individual "
        "de cada caso divergente antes do encerramento do ciclo 2026.")
paragrafo(doc, txt4, sz=10, align="justify"); doc.add_paragraph()
tabela(doc, [
    ["Total com notas em ambos sistemas", f"{total_comparacao:,}", "100%"],
    ["CONVERGENTES (diferenca < 0,01)", f"{total_convergentes:,}", str(pct_convergentes) + "%"],
    ["DIVERGENTES (diferenca >= 0,01)", f"{total_divergentes:,}", str(pct_divergentes) + "%"],
], ["SITUACAO", "QUANTIDADE", "PERCENTUAL"], cor_hdr=COR_AZUL_ESC, widths=[9, 4, 3])
doc.add_paragraph()

if total_divergentes > 0:
    paragrafo(doc, "Amostra das divergencias identificadas (primeiros 20 registros):", bold=True, sz=10)
    div_tab_d = [[
        str(r.get('MATRICULA', '')).strip(),
        str(r.get('NOME SERVIDOR', '')).strip()[:28],
        str(r.get('NOME UNIDADE PRINCIPAL', '')).strip()[:22],
        r.get('NOTA DA AADP', '-'),
        r.get('Nota Final', '-'),
        round(float(r.get('DIFF_NOTA', 0)), 4)
    ] for _, r in divergentes.head(20).iterrows()]
    tabela(doc, div_tab_d,
           ["MATRICULA", "NOME", "UNIDADE", "MAINFRAME", "AADP", "DIFERENCA"],
           cor_hdr=COR_VERM, widths=[2.5, 6, 5, 2.5, 2.5, 2.5])

doc.add_page_break()

# SECAO 5: ANALISE POR RPM
titulo(doc, "5. ANALISE POR REGIAO DE POLICIA MILITAR (RPM)", 2)
paragrafo(doc, "_" * 90, cor="AAAAAA"); doc.add_paragraph()
paragrafo(doc, ("A analise por RPM permite identificar quais regionais apresentam maior ou menor "
                "cobertura de avaliacoes e quais possuem melhores ou piores medias de desempenho, "
                "permitindo direcionar esforcos de regularizacao e monitoramento."), sz=10, align="justify")
doc.add_paragraph()
rpm_top20 = rpm_analise.sort_values('QTD_MILITARES', ascending=False).head(20)
tabela(doc, [
    [str(r['Nome RPM']).strip()[:35], f"{int(r['QTD_MILITARES']):,}",
     f"{int(r['COM_NOTA']):,}", f"{int(r['SEM_NOTA']):,}",
     str(r['PCT_COM_NOTA']) + "%",
     round(float(r['MEDIA_NOTA']), 2) if pd.notna(r['MEDIA_NOTA']) else '-']
    for _, r in rpm_top20.iterrows()
], ["RPM", "TOTAL", "COM NOTA", "SEM NOTA", "% COBERT.", "MEDIA NOTA"],
   widths=[7, 2.5, 2.5, 2.5, 2.5, 2.5])
doc.add_paragraph()

# SECAO 6: SITUACAO FUNCIONAL
titulo(doc, "6. ANALISE POR SITUACAO FUNCIONAL", 2)
paragrafo(doc, "_" * 90, cor="AAAAAA"); doc.add_paragraph()
paragrafo(doc, ("A distribuicao das avaliacoes por situacao funcional evidencia o alcance do "
                "processo avaliativo em relacao as diferentes condicoes de trabalho dos militares."),
          sz=10, align="justify")
doc.add_paragraph()
tabela(doc, [
    [str(r['Sit. Funcional']).strip(), f"{int(r['QTD_MILITARES']):,}",
     f"{int(r['COM_NOTA']):,}", f"{int(r['SEM_NOTA']):,}",
     round(float(r['MEDIA_NOTA']), 2) if pd.notna(r['MEDIA_NOTA']) else '-',
     round(float(r['MIN_NOTA']), 2) if pd.notna(r['MIN_NOTA']) else '-',
     round(float(r['MAX_NOTA']), 2) if pd.notna(r['MAX_NOTA']) else '-']
    for _, r in sit_analise.iterrows()
], ["SIT. FUNCIONAL", "TOTAL", "COM NOTA", "SEM NOTA", "MEDIA", "MINIMA", "MAXIMA"],
   widths=[5, 2, 2, 2, 2, 2, 2])

doc.add_page_break()

# SECAO 7: DISTRIBUICAO DE NOTAS
titulo(doc, "7. DISTRIBUICAO DAS NOTAS POR FAIXA", 2)
paragrafo(doc, "_" * 90, cor="AAAAAA"); doc.add_paragraph()
paragrafo(doc, ("A distribuicao das notas por faixa permite identificar a concentracao do desempenho "
                "dos militares avaliados. Faixas elevadas (9,0 a 10,0) podem indicar inflacao de notas, "
                "enquanto faixas baixas demandam atencao para o desempenho operacional."),
          sz=10, align="justify")
doc.add_paragraph()
dist_geral = df_aadp.groupby('FAIXA_NOTA').size().reset_index(name='QTD')
dist_geral['PCT'] = (dist_geral['QTD'] / dist_geral['QTD'].sum() * 100).round(1)
tabela(doc, [[r['FAIXA_NOTA'], f"{int(r['QTD']):,}", str(r['PCT']) + "%"]
             for _, r in dist_geral.iterrows()],
       ["FAIXA DE NOTA", "QUANTIDADE", "PERCENTUAL"], widths=[8, 4, 4])
doc.add_paragraph()

# SECAO 8: RECOMENDACOES
titulo(doc, "8. RECOMENDACOES AO ALTO COMANDO", 2)
paragrafo(doc, "_" * 90, cor="AAAAAA"); doc.add_paragraph()

recomendacoes = [
    ("Regularizacao urgente das " + f"{total_sem_aadp:,}" + " avaliacoes pendentes",
     "Determinacao de prazo maximo para que as avaliacoes ainda nao registradas no "
     "AADP/SIRHWEB sejam lancadas, sob pena de impacto nos processos de promocao e "
     "remuneracao do ciclo 2026."),
    ("Auditoria das " + f"{total_divergentes:,}" + " divergencias de notas",
     "Instauracao de processo de conferencia e correcao das divergencias identificadas "
     "entre o AADP/SIRHWEB e o Mainframe, garantindo que as notas homologadas reflitam "
     "fidedignamente o desempenho avaliado."),
    ("Monitoramento dos Reconduzidos QPR",
     "Dos " + f"{total_reconduzidos:,}" + " reconduzidos, " + f"{len(recon_sem_nota):,}"
     + " nao possuem avaliacao no AADP/SIRHWEB. Avaliar se o regramento aplicavel exige "
     "inclusao desses militares no ciclo 2026."),
    ("Atencao as RPMs com menor cobertura",
     "As RPMs com percentual de cobertura abaixo de 80% devem receber atencao especial "
     "dos respectivos comandantes regionais para conclusao das avaliacoes pendentes."),
    ("Monitoramento da distribuicao de notas",
     "A analise da distribuicao de notas por faixa deve ser revisada para detectar "
     "possiveis distorcoes ou padroes anomalos que comprometam a fidedignidade do processo."),
    ("Integracao entre AADP/SIRHWEB e Mainframe",
     "Propoe-se estudar a automacao da integracao de notas entre os dois sistemas para "
     "eliminar divergencias futuras e garantir consistencia permanente dos dados."),
]

for i, (tit, desc) in enumerate(recomendacoes):
    p = doc.add_paragraph()
    run = p.add_run(str(i+1) + ". " + tit)
    run.bold = True; run.font.size = Pt(10)
    run.font.color.rgb = rgb(COR_AZUL_ESC); run.font.name = 'Calibri'
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.8)
    run2 = p2.add_run(desc)
    run2.font.size = Pt(10); run2.font.name = 'Calibri'
    doc.add_paragraph()

# RODAPE
doc.add_page_break()
paragrafo(doc, "_" * 90, cor="AAAAAA", align="center"); doc.add_paragraph()
paragrafo(doc, "Belo Horizonte, " + datetime.now().strftime('%d/%m/%Y'), align="center", sz=11)
doc.add_paragraph(); doc.add_paragraph()
paragrafo(doc, "________________________________________", align="center")
paragrafo(doc, "DIRETORIA DE RECURSOS HUMANOS - DRH", bold=True, align="center", sz=11)
paragrafo(doc, "POLICIA MILITAR DE MINAS GERAIS", align="center", sz=10)
doc.add_paragraph()
paragrafo(doc, "Relatorio gerado automaticamente em " + datetime.now().strftime('%d/%m/%Y as %H:%M'),
          sz=8, cor="888888", align="center")

doc.save(WORD_OUT)
print("   Word salvo: " + WORD_OUT)

print("\n" + "=" * 70)
print("  CONCLUIDO COM SUCESSO em " + datetime.now().strftime('%d/%m/%Y %H:%M:%S'))
print("  Arquivos gerados:")
print("  -> " + EXCEL_OUT)
print("  -> " + WORD_OUT)
print("=" * 70)
